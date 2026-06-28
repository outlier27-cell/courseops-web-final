from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ASSET_DIR = Path(os.environ["COURSEOPS_MANUAL_ASSETS"])
OUT_PATH = Path(os.environ["COURSEOPS_MANUAL_OUT"])

INK = RGBColor(0x1F, 0x1F, 0x1F)
MUTED = RGBColor(0x66, 0x66, 0x66)
UIBE_BLUE = RGBColor(0x1F, 0x4E, 0x79)
LIGHT_BLUE = "EAF1F7"
GRID = "B8C7D6"


SCREENSHOTS = [
    ("01-landing-desktop.png", "系统首页", "展示系统定位、项目管理场景和核心功能入口。"),
    ("02-login-desktop.png", "登录页面", "支持学生、教师、管理员三类账号进入系统。"),
    ("03-workspace-desktop.png", "课程项目工作台", "集中展示项目进度、待办材料、风险提示和通知信息。"),
    ("04-projects-desktop.png", "项目中心", "按课程项目聚合查看完成率、风险等级和项目状态。"),
    ("05-project-space-desktop.png", "项目详情", "展示项目基本信息、成员分工、材料清单和关键节点。"),
    ("06-tasks-desktop.png", "材料流转", "管理待提交、待审核、需修改和已通过材料。"),
    ("07-task-detail-desktop.png", "材料详情", "展示材料要求、提交记录、教师反馈和流转历史。"),
    ("10-analytics-fixed-desktop.png", "运营洞察", "通过图表展示完成率、风险分布、提交趋势和健康度指标。"),
    ("09-logs-desktop.png", "审计日志", "记录用户关键操作，体现 MongoDB 日志追踪能力。"),
    ("11-analytics-fixed-mobile.png", "移动端页面", "展示窄屏下的页面适配效果。"),
]


def set_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None, name: str = "宋体") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(paragraph, before: float = 0, after: float = 6, line_spacing: float = 1.25) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_text(doc: Document, text: str, size: float = 11, bold: bool = False, align=None, after: float = 6, color: RGBColor = INK, first_line: bool = False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph_format(p, after=after)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(22)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    set_paragraph_format(p, before=14 if level == 1 else 8, after=6)
    run = p.add_run(text)
    set_font(run, size=15 if level == 1 else 12.5, bold=True, color=UIBE_BLUE if level == 1 else INK, name="黑体")


def add_rule(paragraph, color: str = "1F4E79", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER, fill: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_format(p, after=0, line_spacing=1.2)
    run = p.add_run(text)
    set_font(run, size=11, bold=bold, color=INK, name="宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)


def set_table_borders(table, color: str = GRID) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for name in ("Heading 1", "Heading 2"):
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_cover(doc: Document) -> None:
    add_text(doc, "对外经济贸易大学", size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18, color=UIBE_BLUE, first_line=False)
    add_text(doc, "本科课程期末考核材料", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=72, color=INK)
    add_text(doc, "系统使用说明书", size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=22, color=INK)
    add_text(doc, "题    目：CourseOps 课程项目工作台", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=88, color=INK)

    fields = [
        ("课程名称", "Web 应用程序设计"),
        ("学院", ""),
        ("专业班级", ""),
        ("学号", ""),
        ("姓名", ""),
        ("任课教师", ""),
        ("提交日期", "2026 年 6 月"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        text = f"{label}：{value if value else '____________________________'}"
        run = p.add_run(text)
        set_font(run, size=13, bold=False, color=INK, name="宋体")

    add_text(doc, "说明：提交前请补全学院、专业班级、学号、姓名等个人信息。", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, color=MUTED)
    doc.add_page_break()


def add_metadata_table(doc: Document) -> None:
    table = doc.add_table(rows=5, cols=2)
    table.autofit = False
    set_table_borders(table)
    items = [
        ("前端地址", "http://127.0.0.1:5173/"),
        ("后端地址", "http://127.0.0.1:8080/"),
        ("演示账号", "学生 u-student-001 / 教师 u-teacher-001 / 管理员 u-admin-001，密码均为 123456"),
        ("前端技术", "Vue 3、Vite、TypeScript、Pinia、Arco Design、ECharts"),
        ("后端技术", "Spring Boot 3、MyBatis Plus、MySQL、MongoDB、JWT、Docker；包名 cn.edu.uibe.courseops"),
    ]
    for row, (label, value) in zip(table.rows, items):
        set_cell(row.cells[0], label, bold=True, fill=LIGHT_BLUE)
        set_cell(row.cells[1], value, align=WD_ALIGN_PARAGRAPH.LEFT)


def add_body(doc: Document) -> None:
    add_heading(doc, "一、系统概述", 1)
    add_text(
        doc,
        "CourseOps 课程项目工作台面向课程项目管理场景，覆盖学生提交、教师反馈、材料审核、项目风险跟踪、通知提醒、运营洞察和审计日志等流程。系统将课程项目中分散的材料要求、提交记录和沟通反馈整合到同一个 Web 工作台，便于课堂演示和期末项目验收。",
        first_line=True,
    )
    add_text(
        doc,
        "本说明书按照课程论文/实验报告的正式格式整理，以系统截图为主体，说明主要页面、操作路径和技术实现要点。",
        first_line=True,
    )
    add_text(
        doc,
        "在传统课程项目管理中，学生通常通过聊天软件、邮箱或课堂口头沟通完成材料提交，教师需要手动确认每个学生或每个小组的提交状态，材料版本、反馈内容和修改记录容易分散。CourseOps 的设计目标不是做一个庞大的教务系统，而是围绕“课程项目提交与验收”这一具体场景，提供一套轻量但完整的流程闭环。",
        first_line=True,
    )
    add_text(
        doc,
        "系统重点体现 Web 应用程序设计课程中要求掌握的前后端分离开发、REST 接口设计、数据库持久化、状态管理、可视化展示、权限认证和日志追踪能力。页面风格采用较现代的运营工作台设计，既能满足功能演示，也能体现较好的交互体验。",
        first_line=True,
    )

    add_heading(doc, "二、运行环境与登录信息", 1)
    add_metadata_table(doc)
    add_text(
        doc,
        "运行系统时，需要先启动 Docker 中的 MySQL 和 MongoDB 服务，然后启动 Spring Boot 后端，最后启动 Vite 前端。前端通过统一 API 服务访问后端接口；如果后端暂时不可用，前端会显示本地演示数据提示，保证页面仍可查看，但正式演示时建议使用真实后端数据。",
        first_line=True,
    )
    add_text(
        doc,
        "三个演示账号分别对应学生、教师和管理员视角。学生侧重点是查看项目、提交材料、查看反馈；教师侧重点是审核材料、发布反馈、查看项目风险；管理员侧重点是查看整体数据、审计日志和演示系统状态。",
        first_line=True,
    )

    add_heading(doc, "三、用户角色与业务流程", 1)
    add_heading(doc, "3.1 学生角色", 2)
    add_text(
        doc,
        "学生登录系统后，主要关注自己参与的课程项目、当前需要提交的材料、教师反馈和项目风险提示。学生可以从工作台快速进入待办材料，也可以在项目详情中查看项目成员、材料要求、截止时间和历史提交情况。材料上传后，系统会记录文件名、提交时间、提交人和材料状态，便于后续追踪。",
        first_line=True,
    )
    add_heading(doc, "3.2 教师角色", 2)
    add_text(
        doc,
        "教师角色主要负责查看学生提交情况、审核材料、发布反馈和跟踪项目风险。教师可以将材料状态更新为审核通过，也可以要求学生修改，并通过反馈模块说明需要补充或调整的内容。系统会把教师操作写入日志，使演示时能够体现操作留痕和过程管理。",
        first_line=True,
    )
    add_heading(doc, "3.3 管理员角色", 2)
    add_text(
        doc,
        "管理员角色用于观察系统整体运行情况，包括项目数量、材料提交状态、风险分布、通知和审计日志。管理员视角更适合在演示视频后半部分展示系统完整性，例如说明系统不仅能完成单个材料提交，也能汇总课程项目的整体运营数据。",
        first_line=True,
    )
    add_heading(doc, "3.4 核心业务闭环", 2)
    flow_items = [
        "教师或系统预置课程项目，并配置项目成员、材料要求、截止时间和提交格式。",
        "学生登录后在工作台查看待办材料，进入材料详情页阅读要求并上传文件。",
        "系统保存提交记录，同时更新材料状态、进度和历史流转信息。",
        "教师查看提交内容后进行审核，选择通过或要求修改，并留下具体反馈。",
        "学生根据反馈修改材料，再次提交，直到材料最终通过。",
        "系统把关键操作写入审计日志，并在运营洞察页形成统计图表。",
    ]
    for item in flow_items:
        add_text(doc, item, after=3)

    add_heading(doc, "四、主要功能说明", 1)
    function_sections = [
        ("4.1 登录认证", "登录页面通过后端接口校验用户账号和密码。认证成功后，后端返回 JWT Token 和用户信息，前端将 Token 保存到状态中，并在后续接口请求中携带认证信息。这样可以模拟真实 Web 系统中的登录会话机制。"),
        ("4.2 项目工作台", "工作台是系统登录后的核心入口，展示项目完成率、待处理材料、风险提醒、近期通知和关键统计数据。该页面的作用是让用户一进入系统就能判断当前课程项目是否正常推进，以及下一步最需要处理什么。"),
        ("4.3 项目中心", "项目中心以列表或卡片形式展示课程项目。每个项目包含课程名称、负责教师、完成进度、风险等级、材料数量和成员概况。用户可以从项目中心进入项目详情，查看更细的材料和成员信息。"),
        ("4.4 项目详情", "项目详情页用于解释一个课程项目的完整结构，包括项目简介、成员分工、材料清单、截止时间和风险提示。该页面适合在演示中说明系统如何把课程项目从一个抽象任务拆分为具体材料和具体责任人。"),
        ("4.5 材料流转", "材料流转页面是本系统最核心的业务功能。每份材料都有格式要求、提交状态、负责人、截止日期和进度。学生可以上传材料，教师可以审核状态，系统会记录从待提交到已提交、需修改、已通过的完整变化过程。"),
        ("4.6 教师反馈", "教师反馈用于承接审核意见。教师可以针对某一份材料写出需要修改的内容，学生在材料详情页查看反馈后继续修改。该功能解决了传统沟通中反馈分散、不易查找的问题。"),
        ("4.7 运营洞察", "运营洞察页面通过 ECharts 图表展示课程项目的统计结果，包括材料完成率、风险分布、提交趋势、课程完成度和健康度雷达图。该页面体现了前端数据可视化能力，也让项目看起来更像一个完整的信息系统。"),
        ("4.8 审计日志", "审计日志页面展示用户关键操作，例如上传材料、审核材料、添加反馈、更新状态等。后端使用 MongoDB 保存日志数据，体现了关系型数据库与文档型数据库结合使用的能力。"),
        ("4.9 演示辅助功能", "系统提供退出登录和重置演示数据能力。退出登录用于切换学生、教师、管理员不同角色；重置演示数据用于在课堂演示或录制视频时恢复页面状态，避免多次操作后数据过乱。"),
    ]
    for title, body in function_sections:
        add_heading(doc, title, 2)
        add_text(doc, body, first_line=True)

    add_heading(doc, "五、界面截图与使用说明", 1)
    for index, (filename, title, description) in enumerate(SCREENSHOTS, start=1):
        add_heading(doc, f"5.{index} {title}", 2)
        add_text(doc, description, first_line=True, after=4)
        if title == "系统首页":
            add_text(doc, "首页用于说明系统的整体定位和产品入口。演示时可以先从该页面介绍项目名称、设计目标和系统面向的课程项目管理场景，再点击进入登录页面。首页不承担复杂业务操作，主要起到项目展示和引导作用。", first_line=True, after=5)
        elif title == "登录页面":
            add_text(doc, "登录页面提供真实后端登录入口。演示时建议先使用学生账号登录，完成材料查看和提交流程；随后退出登录并切换教师账号，展示审核与反馈流程；最后可以使用管理员视角查看日志和统计信息。", first_line=True, after=5)
        elif title == "课程项目工作台":
            add_text(doc, "工作台页面适合作为登录后的第一站。用户可以通过顶部导航切换页面，通过统计卡片快速了解项目健康状态，通过待办区域判断哪些材料需要优先处理。该页面承担了系统信息总览的作用。", first_line=True, after=5)
        elif title == "项目中心":
            add_text(doc, "项目中心页面把多个课程项目集中展示，方便用户比较不同项目的进度和风险。对于教师或管理员来说，该页面可以快速判断哪些项目存在延期风险；对于学生来说，可以快速进入自己负责的项目。", first_line=True, after=5)
        elif title == "项目详情":
            add_text(doc, "项目详情页是项目级信息的集中展示区域。页面中包含项目说明、成员分工和材料列表，能够体现系统对项目结构化管理的能力。演示时可以说明每个材料都与项目、负责人、截止时间和状态关联。", first_line=True, after=5)
        elif title == "材料流转":
            add_text(doc, "材料流转页面是提交和审核流程的入口。用户可以按状态查看材料，例如待提交、已提交、需修改和已通过。该页面重点展示材料管理的流程性，而不是简单的文件列表。", first_line=True, after=5)
        elif title == "材料详情":
            add_text(doc, "材料详情页用于展示单份材料的完整上下文，包括格式要求、提交记录、教师反馈和历史流转。演示时可以选择一份材料，说明学生如何根据要求上传文件，教师如何反馈，系统如何保存历史记录。", first_line=True, after=5)
        elif title == "运营洞察":
            add_text(doc, "运营洞察页面体现系统的数据分析能力。页面中的图表不是孤立装饰，而是围绕课程项目管理中的关键问题设计，例如完成率是否足够高、哪些材料缺口明显、哪些项目存在风险。", first_line=True, after=5)
        elif title == "审计日志":
            add_text(doc, "审计日志页面用于证明系统操作是可追溯的。每一条日志都包含操作者、角色、操作内容、目标对象、操作前后状态和来源模块。该页面也是展示 MongoDB 使用价值的重要部分。", first_line=True, after=5)
        elif title == "移动端页面":
            add_text(doc, "移动端截图用于说明页面具备基本响应式适配能力。虽然本项目主要面向桌面端演示，但在较窄屏幕下仍能保持核心内容可读、按钮可点击、布局不严重错位。", first_line=True, after=5)
        image_path = ASSET_DIR / filename
        if image_path.exists():
            width = Inches(2.6) if "mobile" in filename else Inches(6.45)
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=width)
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(caption, after=7)
            caption_run = caption.add_run(f"图 {index}  {title}")
            set_font(caption_run, size=9.5, color=MUTED, name="宋体")

    add_heading(doc, "六、技术实现说明", 1)
    add_heading(doc, "6.1 前端实现", 2)
    add_text(
        doc,
        "前端采用 Vue 3 + Vite + TypeScript 实现，使用组件化方式组织页面。状态管理使用 Pinia，便于在登录状态、项目数据、材料数据和通知信息之间保持统一。页面 UI 采用 Arco Design 组件并结合自定义样式，保证表单、按钮、卡片、导航和图表具有一致的视觉风格。",
        first_line=True,
    )
    add_text(
        doc,
        "前端通过服务层统一封装 API 请求，页面组件不直接拼接后端地址。这样做的好处是接口地址、错误处理、Token 携带和本地演示数据兜底逻辑都可以集中管理，页面代码更清晰，也方便后续维护。",
        first_line=True,
    )
    add_heading(doc, "6.2 后端实现", 2)
    add_text(
        doc,
        "后端采用 Spring Boot 3 构建 REST API，项目包名已经按作业要求设置为 cn.edu.uibe.courseops。后端提供登录认证、项目查询、材料查询、材料上传、状态更新、教师反馈、通知和日志等接口。控制器负责接收请求，服务层处理业务逻辑，Mapper 负责访问 MySQL 数据库。",
        first_line=True,
    )
    add_text(
        doc,
        "认证部分使用 JWT 思路实现。用户登录成功后获得 Token，后续请求可以基于该 Token 识别用户身份。虽然本项目是课程演示系统，但整体结构接近真实企业项目中的前后端分离认证方式。",
        first_line=True,
    )
    add_heading(doc, "6.3 数据库设计", 2)
    add_text(
        doc,
        "系统使用 MySQL 保存结构化业务数据，例如用户、课程、项目、成员、材料、提交记录、教师反馈、通知和材料历史。使用 MongoDB 保存操作日志，因为日志数据结构更灵活，适合按时间顺序追加和查询。这样的设计体现了关系型数据库与文档型数据库的组合使用。",
        first_line=True,
    )
    add_text(
        doc,
        "MySQL 表之间通过项目编号、材料编号、用户编号等字段建立业务关联。材料是系统的核心对象，它连接了项目、提交记录、教师反馈和历史记录。MongoDB 日志则记录关键操作的操作者、角色、目标对象、操作前后状态和时间。",
        first_line=True,
    )
    add_heading(doc, "6.4 包名与作业要求对应", 2)
    add_text(
        doc,
        "作业要求项目包名为 cn.edu.uibe.*。本项目后端入口类、控制器、服务、实体、Mapper、配置类等均位于 cn.edu.uibe.courseops 包路径下，Maven groupId 也调整为 cn.edu.uibe，满足该硬性要求。",
        first_line=True,
    )

    add_heading(doc, "七、典型演示流程", 1)
    steps = [
        "启动 Docker 中的 MySQL 与 MongoDB，再启动 Spring Boot 后端和 Vite 前端。",
        "使用学生账号登录，查看工作台中的待办材料、项目风险和通知信息。",
        "进入材料详情页，按指定格式上传文件，系统生成提交记录和流转历史。",
        "切换教师账号登录，对材料执行审核通过或要求修改，并新增教师反馈。",
        "返回学生视角查看反馈结果，再进入运营洞察和审计日志页面展示系统数据分析与追踪能力。",
    ]
    for i, step in enumerate(steps, start=1):
        add_text(doc, f"{i}. {step}", after=3)
    add_text(
        doc,
        "录制演示视频时，建议不要只点击页面，而是边操作边解释每一步对应的业务意义。例如进入材料详情页时，可以说明该页面解决了“材料要求不清楚、反馈分散、历史版本难追踪”的问题；进入审计日志页时，可以说明系统如何记录关键操作并体现后端 MongoDB 的使用。",
        first_line=True,
    )

    add_heading(doc, "八、提交说明与检查清单", 1)
    add_text(
        doc,
        "正式提交时，请将本说明书与项目源码、数据库配置、演示视频、学生证或电子照片一起提交。演示视频建议包含本人出镜、代码结构说明、运行环境启动、功能操作演示和技术实现讲解。",
        first_line=True,
    )
    checklist = [
        "确认后端源码包名为 cn.edu.uibe.courseops，符合 cn.edu.uibe.* 要求。",
        "确认前端和后端能够正常启动，登录接口可以返回成功结果。",
        "确认学生、教师、管理员三个账号均可用于演示不同角色流程。",
        "确认说明书封面已补全学院、专业班级、学号和姓名。",
        "确认说明书中包含主要页面截图，并能说明每个页面的用途。",
        "确认演示视频包含本人出镜、项目代码介绍、运行环境说明和功能演示。",
        "确认提交压缩包中包含前端、后端、数据库配置、说明书和必要运行说明。",
    ]
    for item in checklist:
        add_text(doc, item, after=3)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(header, after=0)
        run = header.add_run("CourseOps 课程项目工作台系统使用说明书")
        set_font(run, size=9, color=MUTED, name="宋体")
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(footer, after=0)
        f_run = footer.add_run("对外经济贸易大学 Web 应用程序设计期末项目")
        set_font(f_run, size=9, color=MUTED, name="宋体")


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_body(doc)
    add_header_footer(doc)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    main()
